import gradio as gr
import pandas as pd
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from pydantic import BaseModel, Field, field_validator
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.output_parsers import PydanticOutputParser
from docx import Document
from datetime import datetime
import os
import tempfile

# Import modułu hr_assistant (bezwarunkowo)
from hr_assistant import HRAssistant

# Globalna instancja asystenta HR
hr_assistant = None

def initialize_hr_assistant():
    """Inicjalizuje asystenta HR"""
    global hr_assistant
    
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        raise ValueError("Brak klucza OPENAI_API_KEY. Ekspert HR nie może zostać uruchomiony.")
    
    try:
        hr_assistant = HRAssistant(
            openai_api_key=openai_api_key,
            pdf_directory="pdfs"
        )
        print("✅ Asystent HR został zainicjalizowany pomyślnie.")
    except Exception as e:
        print(f"❌ Błąd podczas inicjalizacji asystenta HR: {e}")
        raise

# Inicjalizuj asystenta przy starcie
initialize_hr_assistant()

# --- MODELE DANYCH (PYDANTIC) ---
# Definiują strukturę danych używaną do parsowania odpowiedzi z LLM
# oraz do generowania finalnego wyniku JSON.

class QuestionAnswer(BaseModel):
    """
    Reprezentuje pojedynczą odpowiedź na pytanie analityczne.
    Ten model jest używany przez parser LangChain do strukturyzacji odpowiedzi LLM.
    """
    question_number: int = Field(..., description="Numer pytania z wewnętrznej matrycy.")
    answer: str = Field(..., description="Odpowiedź 'TAK' lub 'NIE'.")
    citation: str = Field(..., description="Cytat z analizowanego tekstu, na podstawie którego udzielono odpowiedzi.")

    @field_validator("answer")
    def validate_answer(cls, v):
        """Walidator sprawdzający, czy odpowiedź to 'TAK' lub 'NIE'."""
        if v not in {"TAK", "NIE"}:
            raise ValueError("Odpowiedź musi być TAK lub NIE")
        return v

class JobAdAnalysis(BaseModel):
    """
    Reprezentuje pełną analizę ogłoszenia, zawierającą listę odpowiedzi.
    Ten model jest używany przez parser LangChain do strukturyzacji odpowiedzi LLM.
    """
    answers: list[QuestionAnswer]

parser = PydanticOutputParser(pydantic_object=JobAdAnalysis)

PROMPT_TEMPLATE_TEXT = """Przeanalizuj poniższe ogłoszenie o pracę pod kątem dostępności dla osób z niepełnosprawnościami.

Ogłoszenie:
{job_ad}

Odpowiedz na następujące pytania:
{questions}

Format odpowiedzi powinien być w następującej strukturze JSON:
{{
  "answers": [
    {{
      "question_number": 1,
      "answer": "TAK/NIE",
      "citation": "dokładny cytat z tekstu"
    }}
  ]
}}
"""

# Wczytanie matrycy danych
matryca_df = pd.read_csv('matryca.csv', header=None,
                         names=['area', 'prompt', 'true', 'false', 'more', 'hint'])

def prepare_questions(df):
    questions_text = ""
    for index, row in df.iterrows():
        question_number = index + 1
        questions_text += f"{question_number} {row['prompt']}\n"
    return questions_text

def doc_to_text(file):
    extension = os.path.splitext(file.name)[1].lower()
    if extension == ".docx":
        loader = Docx2txtLoader(file.name)
    elif extension == ".pdf":
        loader = PyPDFLoader(file.name)
    else:
        return "error"
    pages = loader.load()
    return "\n".join(page.page_content for page in pages)

def is_job_ad(text_fragment: str, model: ChatOpenAI) -> bool:
    """Sprawdza, czy fragment tekstu pochodzi z ogłoszenia o pracę."""
    try:
        prompt = PromptTemplate.from_template(
            "Czy poniższy tekst to fragment ogłoszenia o pracę? Odpowiedz tylko TAK lub NIE.\n\nTekst: {text_to_check}"
        )
        chain = prompt | model | StrOutputParser()
        response = chain.invoke({"text_to_check": text_fragment})
        return "TAK" in response.upper()
    except Exception:
        # W przypadku błędu API, zakładamy, że to nie jest ogłoszenie, aby przerwać przetwarzanie.
        return False

def _generate_report(result: pd.DataFrame, title: str, prefix: str, include_more: bool) -> str:
    """Tworzy dokument Word na podstawie wyników analizy."""
    doc = Document('template.docx')
    doc.add_heading(title, 0)
    doc.add_paragraph(f'Data wygenerowania: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    
    for _, row in result.iterrows():
        doc.add_heading(str(row['area']), 1)
        doc.add_paragraph(str(row['citation']), style='Intense Quote')
        for line in str(row['content']).split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        if include_more and pd.notna(row['more']):
            for line in str(row['more']).split('\n'):
                if line.strip():
                    doc.add_paragraph(line)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename_prefix = f"{prefix}{timestamp}_"
    
    with tempfile.NamedTemporaryFile(delete=False, prefix=filename_prefix, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def create_short_report(result: pd.DataFrame) -> str:
    return _generate_report(
        result,
        title='Raport analizy ogłoszenia o pracę (wersja skrócona)',
        prefix='KoREKtor_short_',
        include_more=False
    )

def create_report(result: pd.DataFrame) -> str:
    return _generate_report(
        result,
        title='Raport analizy ogłoszenia o pracę',
        prefix='KoREKtor_pelny_',
        include_more=True
    )

def process_analysis_results(response, matryca_df):
    """Przetwarza wyniki analizy z LLM i buduje DataFrame."""
    rows = []
    for i, answer_obj in enumerate(response.answers):
        if answer_obj.answer in {"TAK", "NIE"}:
            answer = answer_obj.answer
            # Inwersja odpowiedzi dla pytania nr 10, zgodnie z logiką matrycy.
            if i == 9:
                answer = "NIE" if answer == "TAK" else "TAK"
                
            new_row = {
                'area': matryca_df.area[i],
                'answer': answer,
                'citation': answer_obj.citation,
                'content': matryca_df.true[i] if answer == 'TAK' else matryca_df.false[i],
                'more': matryca_df.more[i]
            }
            rows.append(new_row)
    return pd.DataFrame(rows)

def analyze_job_ad(job_ad, file):
    if not job_ad and not file:
        return None, None, None
    try:
        if file:
            job_ad = doc_to_text(file)
            if job_ad == "error":
                return {"error": "Nieobsługiwany format pliku. Użyj PDF lub DOCX."}, None, None
        
        if not job_ad or job_ad.strip() == "":
            return None, None, None
        
        model = ChatOpenAI(temperature=0, model="gpt-4o-mini")

        # Krok 2: Weryfikacja, czy tekst jest ogłoszeniem o pracę
        text_for_verification = job_ad[:1500]
        if not is_job_ad(text_for_verification, model):
            return {"error": "Przesłany tekst lub plik nie wygląda na ogłoszenie o pracę."}, None, None

        # Krok 3: Główna analiza z użyciem LLM
        questions = prepare_questions(matryca_df)
        prompt_template = PromptTemplate.from_template(PROMPT_TEMPLATE_TEXT)
    
        chain = prompt_template | model | parser
        response = chain.invoke({"job_ad": job_ad, "questions": questions})
    
        # Krok 4: Przetwarzanie odpowiedzi i budowanie DataFrame
        output_df = process_analysis_results(response, matryca_df)
    
        # Krok 5: Generowanie raportów i wyniku JSON
        short_word_file_path = create_short_report(output_df)
        word_file_path = create_report(output_df)
        json_output = output_df.to_dict(orient="records")
        
        return json_output, word_file_path, short_word_file_path
    except Exception as e:
        # Zwracamy błąd w formacie JSON, aby wyświetlić go w interfejsie
        return {"error": f"Wystąpił wewnętrzny błąd serwera: {e}"}, None, None

def ask_hr_assistant(question):
    """Funkcja do zadawania pytań asystentowi HR.

    Args:
        question (str): Pytanie do eksperta HR.

    Returns:
        str: Odpowiedź eksperta HR w formacie Markdown.
    """
    global hr_assistant
    if hr_assistant is None:
        return "⚠️ Ekspert HR nie jest dostępny z powodu błędu inicjalizacji. Sprawdź logi serwera."
    try:
        response = hr_assistant.ask(question)
        answer = f"**Asystent HR:**\n\n{response['answer']}"
        if response.get('sources'):
            answer += f"\n\n📚 **Źródła:**\n"
            for i, source in enumerate(response['sources'][:3], 1):  # Max 3 źródła
                # Rozróżnij między URL a PDF
                if source.get('type') == 'url' and source.get('url'):
                    # Dla URL - stwórz link z tytułem
                    title = source.get('title', 'Strona internetowa')
                    url = source.get('url')
                    source_text = f"{i}. [{title}]({url})"
                    
                    # Dodaj sekcję jeśli istnieje
                    section = source.get('section', '')
                    if section and section != title:
                        source_text += f" - _{section}_"
                        
                else:
                    # Dla PDF - standardowe formatowanie
                    bibliography = source.get('bibliography', source.get('filename', ''))
                    page = source.get('page', '?')
                    section = source.get('section', '')
                    
                    source_text = f"{i}. {bibliography}"
                    if page and page != '?':
                        source_text += f", str. {page}"
                    if section and section != bibliography and section:
                        source_text += f" - _{section}_"
                
                answer += source_text + "\n"
        return answer
    except Exception as e:
        return f"❌ Wystąpił błąd podczas komunikacji z ekspertem HR: {e}"

# --- Interfejs Gradio ---
with gr.Blocks(theme=gr.themes.Soft(primary_hue="blue", secondary_hue="orange"), title="KoREKtor") as demo:
    gr.Image("logo-korektor.png", width=200, show_label=False, container=False, elem_id="logo")

    with gr.Tab("Analizator ogłoszeń"):
        gr.Markdown("## Przeanalizuj ogłoszenie pod kątem dostępności.")
        with gr.Row():
            with gr.Column(scale=1):
                job_ad_input_text = gr.TextArea(label="Wklej treść ogłoszenia", lines=15)
                job_ad_input_file = gr.File(label="lub wgraj plik (PDF/DOCX)", file_types=[".pdf", ".docx"])
                analyze_button = gr.Button("Analizuj", variant="primary")
            with gr.Column(scale=2):
                json_output = gr.JSON(label="Wyniki analizy")
                short_report_output = gr.File(label="Pobierz raport skrócony")
                full_report_output = gr.File(label="Pobierz raport pełny")

        analyze_button.click(
            fn=analyze_job_ad,
            inputs=[job_ad_input_text, job_ad_input_file],
            outputs=[json_output, full_report_output, short_report_output]
        )

    with gr.Tab("Asystent HR"):
        gr.Markdown("## Zadaj pytanie ekspertowi HR.")
        
        question_input = gr.TextArea(label="Pytanie", placeholder="Zadaj pytanie...")
        ask_button = gr.Button("Wyślij", variant="primary")
        answer_output = gr.Markdown(label="Odpowiedź")

        ask_button.click(
            fn=ask_hr_assistant,
            inputs=question_input,
            outputs=answer_output
        )

demo.launch()